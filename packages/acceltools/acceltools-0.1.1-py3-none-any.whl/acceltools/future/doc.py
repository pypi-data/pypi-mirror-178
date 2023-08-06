import os
from decimal import ROUND_HALF_UP, Decimal
from statistics import mean
from sys import exit

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from acceltools.base import ToolBox


def error_temination(error_text):
    input(error_text)
    exit()


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def get_folder_list(parent_directry):
    all_folders = []
    for folder in os.listdir(parent_directry):
        if os.path.isdir(os.path.join(parent_directry, folder)):
            if folder[0:1] is not str("_"):
                all_folders.append(os.path.join(parent_directry, folder))
    all_folders.sort
    return all_folders


def get_file_path(parent_folder, subdir):
    all_files = []
    try:
        for files in os.listdir(os.path.join(parent_folder, subdir)):
            if os.path.isfile(os.path.join(parent_folder, subdir, files)):
                if files[0:1] is not str("_"):
                    all_files.append(os.path.join(parent_folder, subdir, files))
    except FileNotFoundError:
        error_temination("Error on parsing " + os.path.join(parent_folder, subdir))
    all_files.sort()
    if len(all_files) != 1:
        error_temination("Make sure that only one file in the directry: " + os.path.join(parent_folder, subdir))
    return all_files[0]


def initialize_document(document):
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(30)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(12.7)

    document.styles["Normal"].font.name = "Times New Rowman"
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].paragraph_format.space_after = Pt(0)

    return document


def create_xyz(cartesian_list, title, output_path):
    xyz_output = open(output_path, "w")
    output_lines = []
    output_lines.append(str(len(cartesian_list)) + "\n")
    output_lines.append(title + "\n")
    for one_cartesian in cartesian_list:
        output_lines.append("{0[0]:<2} {0[1]:>15} {0[2]:>15} {0[3]:>15}\n".format(one_cartesian))
    output_lines.append("\n")
    xyz_output.writelines(output_lines)
    return None


def get_cartesian(file_path):
    cartesian_file = open(file_path, "r")
    cartesian_file_lines = cartesian_file.readlines()
    cartesian_file.close()

    for i, line in enumerate(cartesian_file_lines):
        if line.find("1\\1\\") > -1:
            archive_line = ""
            after_archive = cartesian_file_lines[i:]
            for line in after_archive:
                archive_line += "".join(line[1:].splitlines())
                if line.find("\\\\@") > -1:
                    break

    archive_entry = archive_line.split("\\\\")
    tmp_cartesian_list = [i.split(",") for i in archive_entry[3].split("\\")[1:]]
    if len(tmp_cartesian_list[0]) == 4:
        cartesian_list = tmp_cartesian_list
    elif len(tmp_cartesian_list[0]) == 5:
        cartesian_list = [[i[0], i[2], i[3], i[4]] for i in tmp_cartesian_list]
    else:
        cartesian_list = None

    center_of_cartesian = []
    for i in range(3):
        center_of_cartesian.append(mean([float(one_cartesian[i + 1]) for one_cartesian in cartesian_list]))
    center_of_cartesian = [round(i, 10) for i in center_of_cartesian]

    for one_cartesian in cartesian_list:
        for i in range(3):
            one_cartesian[i + 1] = str(Decimal(one_cartesian[i + 1]) - Decimal(str(center_of_cartesian[i])))

    return cartesian_list


def get_title(file_path):
    title_file = open(file_path, "r")
    structure_title = title_file.readline()
    title_file.close()
    return structure_title


def format_number(number_string, ROUNDING_DIGITS):
    number_float = float(number_string)

    rounding_string = str("0.")
    for i in range(ROUNDING_DIGITS - 1):
        rounding_string = rounding_string + "0"
    rounding_string = rounding_string + "1"

    formated_string = str(Decimal(str(number_float)).quantize(Decimal(rounding_string), rounding=ROUND_HALF_UP))
    formated_string = formated_string.replace("-", "−")
    return formated_string


def get_scf(file_path):
    spe_file = open(file_path, "r")
    spe_file_lines = spe_file.readlines()
    spe_file.close()
    scf = {}
    scf["scf_file"] = os.path.basename(file_path)

    for i, line in enumerate(spe_file_lines):
        if line.find("SCF Done:") > -1:
            scf["scf_energy"] = spe_file_lines[i].split()[4]
            scf["scf_type"] = spe_file_lines[i].split()[2][2:-1]
    if len(scf) != 3:
        error_temination("Error in getting SCF energy: " + file_path)
    return scf


def get_therm(file_path):
    frq_file = open(file_path, "r")
    frq_file_lines = frq_file.readlines()
    frq_file.close()
    thermal_corrections = {}
    thermal_corrections["therm_file"] = os.path.basename(file_path)

    for i, line in enumerate(frq_file_lines):
        if line.find("Zero-point correction") > -1:
            thermal_corrections["zero_corr"] = line.split()[2]
        if line.find("Thermal correction to Energy") > -1:
            thermal_corrections["to_energy"] = line.split()[4]
        if line.find("Thermal correction to Enthalpy") > -1:
            thermal_corrections["to_enthalpy"] = line.split()[4]
        if line.find("Thermal correction to Gibbs Free Energy") > -1:
            thermal_corrections["to_gibbs"] = line.split()[6]
        if line.find("Harmonic frequencies (cm**-1),") > -1:
            freq_line = frq_file_lines[i + 6]
            thermal_corrections["vib_1"] = freq_line.split()[2]
            thermal_corrections["vib_2"] = freq_line.split()[3]
            thermal_corrections["vib_3"] = freq_line.split()[4]

    if len(thermal_corrections) != 8:
        error_temination("Error in getting thermal correction: " + file_path)

    return thermal_corrections


def describe_one_structure(document, folder_path):
    output_for_summary = {}
    output_for_summary["folder_name"] = os.path.basename(folder_path)
    strucutre_title = get_title(get_file_path(folder_path, "tit"))
    output_for_summary["title"] = strucutre_title

    p_title = document.add_paragraph()
    p_title.add_run("Structure, energy and Cartesian coordinate of ")
    p_title.add_run(strucutre_title).bold = True
    p_title.add_run(":")

    document.add_picture(get_file_path(folder_path, "img"), height=Mm(40))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_img_title = document.add_paragraph()
    p_img_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_title = p_img_title.add_run(strucutre_title)
    img_title.bold = True
    img_title.font.name = "Arial"
    img_title.font.size = Pt(9)

    document.add_paragraph()

    scf = get_scf(get_file_path(folder_path, "spe"))
    thermal_correction = get_therm(get_file_path(folder_path, "frq"))
    output_for_summary.update(scf)
    output_for_summary.update(thermal_correction)

    energy_table = document.add_table(rows=9, cols=2)
    energy_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    energy_table.allow_autofit = False

    custom_widths = (Mm(110), Mm(20))
    for one_row in energy_table.rows:
        one_row.height_rule = WD_ROW_HEIGHT.EXACTLY
        one_row.height = Pt(15)
        for i, one_cell in enumerate(one_row.cells):
            one_cell.width = custom_widths[i]
            one_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in (0, 2, 4):
        energy_table.cell(i, 0).merge(energy_table.cell(i, 1))

    energy_table.cell(0, 0).text = "B3LYP-D3BJ/6-311+G(d,p)-IEFPCM(CH2Cl2)//B3LYP-D3BJ/6-31G(d)-IEFPCM(CH2Cl2):"
    energy_table.cell(1, 0).text = "Gibbs Free Energy (a.u.)   ="
    total_gibbs = float(scf["scf_energy"]) + float(thermal_correction["to_gibbs"])
    output_for_summary["total_gibbs"] = total_gibbs
    energy_table.cell(1, 1).text = format_number(str(total_gibbs), 6)
    energy_table.cell(2, 0).text = "B3LYP-D3BJ/6-311+G(d,p)-IEFPCM(CH2Cl2):"
    energy_table.cell(3, 0).text = "Electronic energy (a.u.)   ="
    energy_table.cell(3, 1).text = format_number(scf["scf_energy"], 6)
    energy_table.cell(4, 0).text = "B3LYP-D3BJ/6-31G(d)-IEFPCM(CH2Cl2):"
    energy_table.cell(5, 0).text = "Zero-point correction (a.u.)   ="
    energy_table.cell(5, 1).text = format_number(thermal_correction["zero_corr"], 6)
    energy_table.cell(6, 0).text = "Thermal correction to Energy (a.u.)   ="
    energy_table.cell(6, 1).text = format_number(thermal_correction["to_energy"], 6)
    energy_table.cell(7, 0).text = "Thermal correction to Enthalpy (a.u.)   ="
    energy_table.cell(7, 1).text = format_number(thermal_correction["to_enthalpy"], 6)
    energy_table.cell(8, 0).text = "Thermal correction to Gibbs Free Energy (a.u.)   ="
    energy_table.cell(8, 1).text = format_number(thermal_correction["to_gibbs"], 6)

    for i in (0, 2, 4, 9):
        if i == len(energy_table.rows):
            for one_cell in energy_table.rows[i - 1].cells:
                set_cell_border(
                    one_cell,
                    bottom={
                        "sz": 5,
                        "val": "single",
                        "color": "#000000",
                    },
                )
        else:
            for one_cell in energy_table.rows[i].cells:
                set_cell_border(
                    one_cell,
                    top={
                        "sz": 5,
                        "val": "single",
                        "color": "#000000",
                    },
                )

    for i in (1, 3, 5, 6, 7, 8):
        energy_table.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph()

    cartesian = get_cartesian(get_file_path(folder_path, "spe"))
    create_xyz(cartesian, strucutre_title, os.path.join(folder_path, strucutre_title + ".xyz"))

    cartesian_table = document.add_table(rows=0, cols=9)
    cartesian_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, one_cartesian in enumerate(cartesian):
        if i % 2 == 0:
            new_row = cartesian_table.add_row()
            new_row.height_rule = WD_ROW_HEIGHT.EXACTLY
            new_row.height = Pt(15)
            row_cells = new_row.cells
            for j in range(9):
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[0].text = one_cartesian[0]
            for j in range(1, 4):
                row_cells[j].text = format_number(one_cartesian[j], 5)
            for j in range(4):
                row_cells[j].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i % 2 == 1:
            row_cells[5].text = one_cartesian[0]
            for j in range(1, 4):
                row_cells[j + 5].text = format_number(one_cartesian[j], 5)
            for j in range(4):
                row_cells[j + 5].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_page_break()
    return output_for_summary


def export_summary(summary, current_path):
    output_lines = [
        "Folder Name,Title,SCF File,SCF,SCF type,Freq File,Zero-point correction,"
        "Thermal correction to Energy,Thermal correction to Enthalpy,"
        "Thermal correction to Gibbs Free Energy,"
        "Frequency 1,Frequency 2,Frequency 3,"
        "Gibbs Free Energy (Hartree),Gibbs Free Energy (kcal/mol)"
        ",\n"
    ]
    for one_summary in summary:
        one_line = []
        one_line.append(one_summary["folder_name"])
        one_line.append(one_summary["title"])
        one_line.append(one_summary["scf_file"])
        one_line.append(one_summary["scf_energy"])
        one_line.append(one_summary["scf_type"])
        one_line.append(one_summary["therm_file"])
        one_line.append(one_summary["zero_corr"])
        one_line.append(one_summary["to_energy"])
        one_line.append(one_summary["to_enthalpy"])
        one_line.append(one_summary["to_gibbs"])
        one_line.append(one_summary["vib_1"])
        one_line.append(one_summary["vib_2"])
        one_line.append(one_summary["vib_3"])
        one_line.append(str(one_summary["total_gibbs"]))
        one_line.append(str(one_summary["total_gibbs"] * 627.5095))
        one_line_str = ""
        for one_item in one_line:
            one_line_str = one_line_str + one_item + ","
        output_lines.append(one_line_str + "\n")

    summary_csv = open(os.path.join(current_path, "summary.csv"), "w")
    summary_csv.writelines(output_lines)
    summary_csv.close
    return None


def main():
    current_path = os.path.dirname(os.path.abspath(__file__))
    folder_list = get_folder_list(current_path)

    document = Document()
    initialize_document(document)

    summary = []
    for folder_path in folder_list:
        summary.append(
            describe_one_structure(
                document,
                folder_path,
            )
        )
        print(os.path.basename(folder_path) + " was processed.")

    print("exporting summary.csv")
    export_summary(summary, current_path)

    print("exporting si.docx")
    try:
        document.save("si.docx")
    except PermissionError:
        error_temination("Permission Error on saving word file")


class DocBox(ToolBox):
    pass
